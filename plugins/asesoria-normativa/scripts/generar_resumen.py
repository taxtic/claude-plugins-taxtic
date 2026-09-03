"""Construcción del .docx del resumen normativo con marca TAXTIC.

El título, la bajada y la fecha se ensamblan desde fuente.json; los títulos de
sección y los subtítulos salen del catálogo. El modelo no aporta ninguno de esos
textos: por eso no necesitan respaldo de cita.
"""
import importlib.util as _il, os as _os

def _cargar_vecino(nombre):
    ruta = _os.path.join(_os.path.dirname(__file__), nombre + ".py")
    spec = _il.spec_from_file_location(nombre, ruta)
    modulo = _il.module_from_spec(spec)
    spec.loader.exec_module(modulo)
    return modulo

catalogo = _cargar_vecino("catalogo")

NARANJA = "D57A23"
GRIS_FONDO = "F7F7F7"
GRIS_TEXTO = "5A5A5A"
BLANCO = "FFFFFF"
FUENTE_TIPOGRAFICA = "Arial"

ETIQUETAS_DE_CALLOUT = {
    "novedad": "NOVEDAD",
    "critico": "CRÍTICO",
    "proteccion": "PROTECCIÓN AL CONTRIBUYENTE",
}
NOMBRES_DE_TIPO = {"circular": "Circular", "resolucion": "Resolución Exenta",
                   "oficio": "Oficio"}
LEYENDA_IA = ("Este documento ha sido generado con apoyo de inteligencia artificial a "
              "partir de la información proporcionada. Debe ser revisado y validado por "
              "un profesional tributario antes de su uso.")
_ASSETS = _os.path.join(_os.path.dirname(__file__), "..", "assets")


def titulo_del_documento(fuente):
    """Identidad del documento, ensamblada sin intervención del modelo."""
    nombre = NOMBRES_DE_TIPO.get(fuente["tipo"], fuente["tipo"].capitalize())
    return f"{nombre} N° {fuente['numero']} de {fuente['anio']} del SII"


def _sombrear(elemento, color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), color)
    elemento.append(sombra)


def _borde_izquierdo(parrafo, color, grosor=24):
    """Barra vertical naranja del callout: el borde del cuadrado de marca."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    propiedades = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    izquierdo = OxmlElement("w:left")
    izquierdo.set(qn("w:val"), "single")
    izquierdo.set(qn("w:sz"), str(grosor))
    izquierdo.set(qn("w:space"), "8")
    izquierdo.set(qn("w:color"), color)
    bordes.append(izquierdo)
    propiedades.append(bordes)
    _sombrear(propiedades, GRIS_FONDO)


def _campo_pagina(parrafo):
    """Número de página dinámico."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    corrida = parrafo.add_run()
    inicio = OxmlElement("w:fldChar"); inicio.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve"); instruccion.text = "PAGE"
    fin = OxmlElement("w:fldChar"); fin.set(qn("w:fldCharType"), "end")
    corrida._r.append(inicio); corrida._r.append(instruccion); corrida._r.append(fin)


def _escribir(parrafo, texto, *, tamano=10.5, negrita=False, cursiva=False,
              color=None, mayuscula=False):
    from docx.shared import Pt, RGBColor
    corrida = parrafo.add_run(texto.upper() if mayuscula else texto)
    corrida.font.name = FUENTE_TIPOGRAFICA
    corrida.font.size = Pt(tamano)
    corrida.bold = negrita
    corrida.italic = cursiva
    if color:
        corrida.font.color.rgb = RGBColor.from_string(color)
    return corrida


def _agregar_parrafo(documento, texto, **estilo):
    parrafo = documento.add_paragraph()
    _escribir(parrafo, texto, **estilo)
    return parrafo


def _agregar_lista(documento, items):
    # Lista nativa de Word, no un carácter de viñeta tecleado: sobrevive a que el
    # usuario edite, reordene o continúe la lista en el documento entregado.
    from docx.shared import Pt
    for item in items:
        parrafo = documento.add_paragraph(style="List Bullet")
        parrafo.paragraph_format.space_after = Pt(2)
        _escribir(parrafo, item["texto"])


def _agregar_callout(documento, bloque):
    from docx.shared import Pt
    etiqueta = documento.add_paragraph()
    etiqueta.paragraph_format.space_before = Pt(8)
    etiqueta.paragraph_format.space_after = Pt(0)
    _borde_izquierdo(etiqueta, NARANJA)
    _escribir(etiqueta, ETIQUETAS_DE_CALLOUT[bloque["variante"]],
              tamano=8, negrita=True, color=NARANJA)
    cuerpo = documento.add_paragraph()
    cuerpo.paragraph_format.space_after = Pt(8)
    _borde_izquierdo(cuerpo, NARANJA)
    _escribir(cuerpo, bloque["texto"])


def _agregar_tabla(documento, bloque):
    from docx.shared import Pt
    encabezado, filas = bloque["encabezado"], bloque["filas"]
    tabla = documento.add_table(rows=1 + len(filas), cols=len(encabezado))
    tabla.style = "Table Grid"
    for c, celda in enumerate(encabezado):
        casilla = tabla.rows[0].cells[c]
        casilla.text = ""
        _sombrear(casilla._tc.get_or_add_tcPr(), NARANJA)
        _escribir(casilla.paragraphs[0], celda["texto"], tamano=9.5,
                  negrita=True, color=BLANCO)
    for f, fila in enumerate(filas, start=1):
        for c, celda in enumerate(fila):
            casilla = tabla.rows[f].cells[c]
            casilla.text = ""
            if f % 2 == 0:
                _sombrear(casilla._tc.get_or_add_tcPr(), GRIS_FONDO)
            _escribir(casilla.paragraphs[0], celda["texto"], tamano=9.5, negrita=(c == 0))


def _agregar_bloque(documento, bloque):
    from docx.shared import Pt
    tipo = bloque["tipo"]
    if tipo == "subtitulo":
        parrafo = documento.add_paragraph()
        parrafo.paragraph_format.space_before = Pt(8)
        parrafo.paragraph_format.space_after = Pt(2)
        _escribir(parrafo, catalogo.SUBTITULOS[bloque["rotulo_id"]],
                  tamano=10.5, negrita=True)
    elif tipo == "parrafo":
        _agregar_parrafo(documento, bloque["texto"])
    elif tipo == "nota":
        _agregar_parrafo(documento, bloque["texto"], tamano=8.5, cursiva=True,
                         color=GRIS_TEXTO)
    elif tipo == "lista":
        _agregar_lista(documento, bloque["items"])
    elif tipo == "callout":
        _agregar_callout(documento, bloque)
    elif tipo == "tabla":
        _agregar_tabla(documento, bloque)


def _agregar_encabezado_y_pie(documento):
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    seccion = documento.sections[0]
    imagotipo = _os.path.join(_ASSETS, "imagotipo-principal-negro.png")
    if _os.path.isfile(imagotipo):
        parrafo = seccion.header.paragraphs[0]
        parrafo.add_run().add_picture(imagotipo, height=Inches(0.28))
    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    isologo = _os.path.join(_ASSETS, "isologo-naranjo.png")
    if _os.path.isfile(isologo):
        pie.add_run().add_picture(isologo, height=Inches(0.16))
    _escribir(pie, "  TAXTIC Asesoría Tributaria Integral  ·  ",
              tamano=8, color=GRIS_TEXTO)
    _campo_pagina(pie)


def _ordenar_secciones(secciones, tipo):
    """Orden de emisión según el perfil, no el orden en que vinieron escritas.

    El orden lo define catalogo.orden_de_emision() y no se duplica acá. El
    ordenamiento es estable, así que varias secciones del mismo id —las
    `materia` repetibles— conservan su orden relativo.
    """
    orden = catalogo.orden_de_emision(tipo)
    posicion = {id_seccion: i for i, id_seccion in enumerate(orden)}
    return sorted(secciones, key=lambda s: posicion.get(s["id"], len(orden)))


def construir_docx(fuente, resumen, salida):
    from docx import Document
    from docx.shared import Pt, Cm
    documento = Document()
    normal = documento.styles["Normal"].font
    normal.name = FUENTE_TIPOGRAFICA
    normal.size = Pt(10.5)
    for seccion in documento.sections:
        seccion.page_width, seccion.page_height = Cm(21), Cm(29.7)
        for margen in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(seccion, margen, Cm(2))

    _agregar_encabezado_y_pie(documento)

    encabezado = documento.add_paragraph()
    encabezado.paragraph_format.space_after = Pt(0)
    _escribir(encabezado, "RESUMEN EJECUTIVO", tamano=14, negrita=True, color=NARANJA)

    titulo = documento.add_paragraph()
    titulo.paragraph_format.space_after = Pt(0)
    _escribir(titulo, titulo_del_documento(fuente), tamano=12.5, negrita=True)

    if fuente.get("materia"):
        bajada = documento.add_paragraph()
        bajada.paragraph_format.space_after = Pt(0)
        _escribir(bajada, fuente["materia"], tamano=9.5, cursiva=True, color=GRIS_TEXTO)

    partes = []
    if fuente.get("fecha_documento"):
        partes.append(f"Fecha del documento: {fuente['fecha_documento']}")
    if resumen.get("meta", {}).get("elaborado_por"):
        partes.append(f"Elaborado por {resumen['meta']['elaborado_por']}")
    partes.append("TAXTIC Asesoría Tributaria Integral")
    metadata = documento.add_paragraph()
    metadata.paragraph_format.space_after = Pt(10)
    _escribir(metadata, " · ".join(partes), tamano=8, color=GRIS_TEXTO)

    for numero, seccion in enumerate(
            _ordenar_secciones(resumen["secciones"], fuente["tipo"]), start=1):
        titulo_seccion = (seccion["titulo"]["texto"] if seccion["id"] == "materia"
                          else catalogo.titulo_de(seccion["id"], fuente["tipo"]))
        parrafo = documento.add_paragraph()
        parrafo.paragraph_format.space_before = Pt(12)
        parrafo.paragraph_format.space_after = Pt(4)
        _escribir(parrafo, "■ ", tamano=9, color=NARANJA)
        _escribir(parrafo, f"{numero}. {titulo_seccion}", tamano=11.5,
                  negrita=True, color=NARANJA)
        for bloque in seccion["bloques"]:
            _agregar_bloque(documento, bloque)

    leyenda = documento.add_paragraph()
    leyenda.paragraph_format.space_before = Pt(16)
    _borde_izquierdo(leyenda, GRIS_TEXTO, grosor=12)
    _escribir(leyenda, LEYENDA_IA, tamano=8, cursiva=True, color=GRIS_TEXTO)

    documento.save(salida)
    return salida


def _main():
    import argparse, json
    analizador = argparse.ArgumentParser(description="Arma el .docx del resumen normativo")
    analizador.add_argument("fuente")
    analizador.add_argument("resumen")
    analizador.add_argument("--docx", default="resumen.docx")
    argumentos = analizador.parse_args()
    with open(argumentos.fuente, encoding="utf-8") as archivo:
        fuente = json.load(archivo)
    with open(argumentos.resumen, encoding="utf-8") as archivo:
        resumen = json.load(archivo)
    construir_docx(fuente, resumen, argumentos.docx)
    print(f"Generado {argumentos.docx}")


if __name__ == "__main__":
    _main()
