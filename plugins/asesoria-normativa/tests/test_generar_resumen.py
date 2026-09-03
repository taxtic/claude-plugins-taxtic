import importlib.util, os, pytest

docx = pytest.importorskip("docx")

def _cargar(nombre):
    ruta = os.path.join(os.path.dirname(__file__), "..", "scripts", nombre + ".py")
    spec = importlib.util.spec_from_file_location(nombre, ruta)
    modulo = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(modulo)
    return modulo

gr = _cargar("generar_resumen")

FUENTE = {
    "tipo": "circular", "numero": "35", "anio": 2026,
    "fecha_documento": "31 de agosto de 2026",
    "materia": "Actualiza instrucciones sobre mecanismos de impugnación administrativa.",
    "metricas": {"paginas": 2, "caracteres": 400},
}

RESUMEN = {
    "meta": {"elaborado_por": "Asesor Tributario"},
    "secciones": [
        {"id": "tema", "bloques": [
            {"tipo": "parrafo", "afirmacion": "citada", "texto": "Sistematiza los mecanismos.",
             "cita": "a" * 45, "pagina": 1}]},
        {"id": "materia",
         "titulo": {"texto": "RAV (art. 123 bis CT)", "cita": "b" * 45, "pagina": 1},
         "bloques": [
             {"tipo": "subtitulo", "rotulo_id": "silencio"},
             {"tipo": "callout", "variante": "novedad", "afirmacion": "citada",
              "texto": "Ahora exige individualizar cada vicio.", "cita": "c" * 45, "pagina": 2},
             {"tipo": "tabla", "afirmacion": "citada",
              "encabezado": [{"texto": ""}, {"texto": "RAV", "cita": "d" * 45, "pagina": 1}],
              "filas": [[{"texto": "Plazo", "cita": "e" * 45, "pagina": 1},
                         {"texto": "30 días", "cita": "f" * 45, "pagina": 1}]]}]},
        {"id": "gestiones", "bloques": [
            {"tipo": "lista", "afirmacion": "derivada",
             "items": [{"texto": "Revisar los expedientes en curso."}]}]},
        {"id": "a_verificar", "bloques": [
            {"tipo": "lista", "afirmacion": "derivada",
             "items": [{"texto": "Confirmar el Diario Oficial."}]}]},
    ],
}


def _textos(documento):
    partes = [p.text for p in documento.paragraphs]
    for tabla in documento.tables:
        for fila in tabla.rows:
            partes += [c.text for c in fila.cells]
    return partes


def test_titulo_se_ensambla_desde_la_fuente():
    assert gr.titulo_del_documento(FUENTE) == "Circular N° 35 de 2026 del SII"

def test_bajada_es_la_materia_de_la_fuente(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    textos = _textos(docx.Document(salida))
    assert any(FUENTE["materia"] in t for t in textos)

def test_encabezado_lleva_resumen_ejecutivo_en_mayuscula(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    assert "RESUMEN EJECUTIVO" in _textos(docx.Document(salida))

def test_titulos_de_seccion_salen_del_catalogo(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    textos = _textos(docx.Document(salida))
    assert any("Tema central" in t for t in textos)
    assert any("Gestiones a considerar" in t for t in textos)
    assert any("Puntos a confirmar" in t for t in textos)

def test_titulo_de_materia_sale_del_resumen(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    assert any("RAV (art. 123 bis CT)" in t for t in _textos(docx.Document(salida)))

def test_subtitulo_se_resuelve_por_rotulo_id(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    assert any("Resolución y silencio administrativo" in t
               for t in _textos(docx.Document(salida)))

def test_callout_lleva_su_etiqueta_en_mayuscula(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    assert any("NOVEDAD" in t for t in _textos(docx.Document(salida)))

def test_tabla_se_renderiza_con_sus_celdas(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    documento = docx.Document(salida)
    assert len(documento.tables) == 1
    assert documento.tables[0].rows[1].cells[1].text == "30 días"

def test_leyenda_ia_presente(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    assert any("inteligencia artificial" in t for t in _textos(docx.Document(salida)))

def test_elaborado_por_en_la_metadata(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    assert any("Asesor Tributario" in t for t in _textos(docx.Document(salida)))

def test_sin_elaborado_por_la_linea_se_omite(tmp_path):
    import copy
    resumen = copy.deepcopy(RESUMEN)
    resumen["meta"] = {}
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, resumen, salida)
    textos = " ".join(_textos(docx.Document(salida)))
    assert "Elaborado por" not in textos

def _titulos_de_seccion(documento):
    """Los títulos numerados, en el orden en que quedaron en el documento."""
    import re
    encontrados = []
    for parrafo in documento.paragraphs:
        encontrado = re.match(r"^■\s*\d+\.\s*(.+)$", parrafo.text.strip())
        if encontrado:
            encontrados.append(encontrado.group(1))
    return encontrados


def test_las_secciones_se_ordenan_segun_el_catalogo(tmp_path):
    """El modelo puede escribirlas en cualquier orden; el builder las acomoda."""
    import copy
    resumen = copy.deepcopy(RESUMEN)
    resumen["secciones"] = [resumen["secciones"][3],   # a_verificar
                            resumen["secciones"][2],   # gestiones
                            resumen["secciones"][1],   # materia
                            resumen["secciones"][0]]   # tema
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, resumen, salida)
    titulos = _titulos_de_seccion(docx.Document(salida))
    assert titulos[0] == "Tema central"
    assert titulos[-2:] == ["Gestiones a considerar", "Puntos a confirmar"]

def test_las_materias_repetidas_conservan_su_orden_relativo(tmp_path):
    import copy
    resumen = copy.deepcopy(RESUMEN)
    segunda = copy.deepcopy(resumen["secciones"][1])
    segunda["titulo"] = {"texto": "RAF (art. 6 letra B CT)", "cita": "g" * 45, "pagina": 2}
    # se insertan desordenadas respecto del resto, pero entre sí en este orden
    resumen["secciones"] = [resumen["secciones"][2], resumen["secciones"][1],
                            segunda, resumen["secciones"][3], resumen["secciones"][0]]
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, resumen, salida)
    titulos = _titulos_de_seccion(docx.Document(salida))
    assert titulos.index("RAV (art. 123 bis CT)") < titulos.index("RAF (art. 6 letra B CT)")
    assert titulos[-2:] == ["Gestiones a considerar", "Puntos a confirmar"]

def test_el_orden_no_esta_duplicado_en_el_builder():
    """catalogo.py es la única fuente del orden."""
    ruta = os.path.join(os.path.dirname(__file__), "..", "scripts", "generar_resumen.py")
    with open(ruta, encoding="utf-8") as archivo:
        codigo = archivo.read()
    assert "orden_de_emision" in codigo
    assert "a_verificar" not in codigo

def test_naranja_de_marca_en_el_titulo(tmp_path):
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    documento = docx.Document(salida)
    encabezado = next(p for p in documento.paragraphs if p.text == "RESUMEN EJECUTIVO")
    assert str(encabezado.runs[0].font.color.rgb) == "D57A23"


def test_la_vineta_es_un_cuadrado_naranja(tmp_path):
    """El cuadrado es el símbolo de marca; el estilo nativo trae un punto negro."""
    from docx.oxml.ns import qn
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, RESUMEN, salida)
    numeracion = docx.Document(salida).part.numbering_part.element
    glifos, colores = set(), set()
    for definicion in numeracion.findall(qn("w:abstractNum")):
        nivel = definicion.find(qn("w:lvl"))
        if nivel is None:
            continue
        texto = nivel.find(qn("w:lvlText"))
        fuente = nivel.find(qn("w:rPr"))
        if texto is not None:
            glifos.add(texto.get(qn("w:val")))
        if fuente is not None and fuente.find(qn("w:color")) is not None:
            colores.add(fuente.find(qn("w:color")).get(qn("w:val")))
    assert "▪" in glifos
    assert gr.NARANJA in colores

def test_una_fuente_sin_identidad_no_produce_un_titulo_falso(tmp_path):
    """Sin esto el título salía "Circular N° 35 de None del SII"."""
    import copy, pytest
    for campo in ("tipo", "numero", "anio"):
        fuente = copy.deepcopy(FUENTE)
        fuente[campo] = None
        with pytest.raises(gr.FuenteIncompleta) as e:
            gr.titulo_del_documento(fuente)
        assert campo in str(e.value)

def test_un_bloque_desconocido_no_desaparece_en_silencio(tmp_path):
    import copy, pytest
    resumen = copy.deepcopy(RESUMEN)
    resumen["secciones"][0]["bloques"][0] = {"tipo": "grafico", "texto": "x"}
    with pytest.raises(ValueError) as e:
        gr.construir_docx(FUENTE, resumen, str(tmp_path / "r.docx"))
    assert "grafico" in str(e.value)

def test_una_seccion_no_admitida_no_se_emite_despues_del_cierre(tmp_path):
    """Caía al final del orden, o sea después de gestiones y a_verificar."""
    import copy, pytest
    resumen = copy.deepcopy(RESUMEN)
    resumen["secciones"].append({"id": "caso_consultado", "bloques": [
        {"tipo": "parrafo", "afirmacion": "citada", "texto": "x",
         "cita": "a" * 45, "pagina": 1}]})
    with pytest.raises(ValueError) as e:
        gr.construir_docx(FUENTE, resumen, str(tmp_path / "r.docx"))
    assert "caso_consultado" in str(e.value)


# --- Integración gate → builder. El producto es una cadena, y ningún test
# --- verificaba que lo que pasa el gate sea construible.

def _fuente_y_resumen_reales():
    """Una fuente coherente consigo misma y un resumen que la cita de verdad."""
    import importlib.util, os
    ruta = os.path.join(os.path.dirname(__file__), "..", "scripts", "extraer_fuente.py")
    spec = importlib.util.spec_from_file_location("extraer_fuente", ruta)
    ef = importlib.util.module_from_spec(spec); spec.loader.exec_module(ef)
    pagina = ("El Servicio deberá pronunciarse dentro del plazo de noventa (90) días "
              "hábiles administrativos contados desde la presentación del recurso.")
    fuente = ef.construir_fuente(
        [pagina], {"clase": "pdf", "ruta": "x.pdf"},
        {"tipo": "circular", "numero": "35", "materia": "Instruye sobre plazos",
         "fecha_documento": "31 de agosto de 2026"})
    resumen = {
        "meta": {"elaborado_por": "Asesor Tributario"},
        "secciones": [
            {"id": "tema", "bloques": [
                {"tipo": "parrafo", "afirmacion": "citada",
                 "texto": "El SII debe pronunciarse dentro de 90 días hábiles administrativos.",
                 "cita": "deberá pronunciarse dentro del plazo de noventa (90) días hábiles administrativos",
                 "pagina": 1}]},
            {"id": "gestiones", "bloques": [
                {"tipo": "lista", "afirmacion": "derivada",
                 "items": [{"texto": "Revisar los expedientes en curso."}]}]},
            {"id": "a_verificar", "bloques": [
                {"tipo": "lista", "afirmacion": "derivada",
                 "items": [{"texto": "Confirmar la publicación en el Diario Oficial."}]}]},
        ],
    }
    return ef, fuente, resumen


def test_lo_que_pasa_el_gate_se_puede_construir(tmp_path):
    import importlib.util, os
    ruta = os.path.join(os.path.dirname(__file__), "..", "scripts", "verificar_citas.py")
    spec = importlib.util.spec_from_file_location("verificar_citas", ruta)
    vc = importlib.util.module_from_spec(spec); spec.loader.exec_module(vc)
    _, fuente, resumen = _fuente_y_resumen_reales()
    filas = vc.verificar(resumen, fuente)          # no levanta
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(fuente, resumen, salida)     # tampoco
    textos = _textos(docx.Document(salida))
    assert any("90 días hábiles" in t for t in textos)
    assert filas

def test_el_builder_no_pisa_un_documento_existente(tmp_path):
    existente = tmp_path / "resumen.docx"
    existente.write_bytes(b"version anterior del contador")
    assert gr._ruta_sin_pisar(str(existente)) == str(tmp_path / "resumen-2.docx")
    assert existente.read_bytes() == b"version anterior del contador"


def test_el_bloque_nota_se_renderiza(tmp_path):
    """Toda su ruta estaba sin cubrir: existe en el catálogo, el esquema y el
    builder, y ningún test lo tocaba."""
    import copy
    resumen = copy.deepcopy(RESUMEN)
    resumen["secciones"][1]["bloques"].append({
        "tipo": "nota", "afirmacion": "citada",
        "texto": "CT = Código Tributario.", "cita": "n" * 45, "pagina": 1})
    salida = str(tmp_path / "r.docx")
    gr.construir_docx(FUENTE, resumen, salida)
    assert any("CT = Código Tributario." in t for t in _textos(docx.Document(salida)))
